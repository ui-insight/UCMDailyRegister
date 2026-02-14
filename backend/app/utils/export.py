"""Word document export for newsletters using python-docx."""

import io
import re
from datetime import date

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def export_newsletter_docx(
    newsletter_type: str,
    publish_date: date,
    sections: list[dict],
) -> io.BytesIO:
    """Generate a Word document for a newsletter.

    Args:
        newsletter_type: "tdr" or "myui"
        publish_date: The publish date
        sections: List of section dicts, each with:
            - name: Section name
            - items: List of item dicts with 'final_headline', 'final_body'

    Returns:
        BytesIO buffer containing the .docx file
    """
    doc = Document()

    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    title_text = (
        "The Daily Register" if newsletter_type == "tdr" else "My UI"
    )
    title = doc.add_heading(title_text, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0xB5, 0x8A, 0x2A)  # Gold

    # Date line
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(publish_date.strftime("%A, %B %d, %Y"))
    date_run.font.size = Pt(12)
    date_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()  # Spacer

    for section in sections:
        if not section.get("items"):
            continue

        # Section heading
        heading = doc.add_heading(section["name"], level=1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0x8C, 0x6E, 0x1A)  # Dark gold

        # Items
        for item in section["items"]:
            headline = item["final_headline"]
            body = item["final_body"]

            # Headline as bold paragraph
            h_para = doc.add_paragraph()
            h_run = h_para.add_run(headline)
            h_run.bold = True
            h_run.font.size = Pt(11)

            # Body — parse HTML anchor tags into hyperlinks
            _add_body_with_links(doc, body)

            doc.add_paragraph()  # Spacer between items

    # Footer
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(
        "University of Idaho | University Communications and Marketing"
    )
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _add_body_with_links(doc: Document, body: str) -> None:
    """Add body text to the document, converting HTML <a> tags to styled text.

    Since python-docx doesn't natively support hyperlinks easily,
    we render link text as blue underlined text with the URL in parentheses.
    """
    # Pattern to match <a href="url">text</a>
    link_pattern = re.compile(r'<a\s+href=["\']([^"\']+)["\']>([^<]+)</a>', re.IGNORECASE)

    para = doc.add_paragraph()
    last_end = 0

    for match in link_pattern.finditer(body):
        # Add text before the link
        if match.start() > last_end:
            before_text = body[last_end:match.start()]
            run = para.add_run(before_text)
            run.font.size = Pt(11)

        # Add link text (blue, underlined)
        url = match.group(1)
        link_text = match.group(2)
        link_run = para.add_run(link_text)
        link_run.font.size = Pt(11)
        link_run.font.color.rgb = RGBColor(0x00, 0x56, 0xB3)
        link_run.font.underline = True

        last_end = match.end()

    # Add remaining text after last link
    if last_end < len(body):
        remaining = body[last_end:]
        run = para.add_run(remaining)
        run.font.size = Pt(11)
